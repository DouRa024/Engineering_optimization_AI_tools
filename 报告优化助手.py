import datetime
import os
import streamlit as st
import requests
from io import BytesIO
from docx import Document
from PyPDF2 import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from theme_manager import apply_theme
import time
apply_theme()
st.set_page_config(page_title="报告优化助手", page_icon="🎯", layout="wide", initial_sidebar_state="expanded")

# 初始化历史记录和确认的上传文件存储
if "history" not in st.session_state:
    st.session_state.history = []
if "confirmed_uploaded_file" not in st.session_state:
    st.session_state.confirmed_uploaded_file = None
if "role_select_1" not in st.session_state:
    st.session_state.role_select_1 = "无角色（原文分析）"
if "role_switching_1" not in st.session_state:
    st.session_state.role_switching_1 = False
if "role_switch_info_1" not in st.session_state:
    st.session_state.role_switch_info_1 = ""

if "user_direction" not in st.session_state:
    st.session_state.user_direction = ""

# --------- DeepSeek API 调用 ---------
def call_deepseek_api(prompt: str) -> str:
    api_key = os.getenv("DEEPSEEK_API_KEY")
    if not api_key:
        st.error("请先设置环境变量 DEEPSEEK_API_KEY")
        return ""

    url = "https://api.deepseek.com/beta/v1/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    data = {
        "model": "deepseek-chat",
        "prompt": prompt,
        "max_tokens": 1500,
        "temperature": 0.7
    }

    resp = requests.post(url, headers=headers, json=data)
    if resp.status_code != 200:
        st.error(f"API调用失败: {resp.status_code} - {resp.text}")
        return ""

    result = resp.json()
    return result.get("choices", [{}])[0].get("text", "").strip()

# --------- 文件文本提取 ---------
def extract_text_from_pdf(file) -> str:
    pdf = PdfReader(file)
    text = [page.extract_text() for page in pdf.pages if page.extract_text()]
    return "\n".join(text)

def extract_text_from_docx(file) -> str:
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_txt(file) -> str:
    return file.read().decode("utf-8")

# --------- 添加优化内容 ---------
def append_text_to_docx(original_file, appended_text):
    doc = Document(original_file)
    doc.add_page_break()
    doc.add_paragraph("=== 优化内容 ===")
    doc.add_paragraph(appended_text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def append_text_to_pdf(original_file, appended_text):
    reader = PdfReader(original_file)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)

    packet = BytesIO()
    can = canvas.Canvas(packet, pagesize=letter)
    text_object = can.beginText(40, 750)
    for line in ("=== 优化内容 ===\n" + appended_text).split("\n"):
        text_object.textLine(line)
    can.drawText(text_object)
    can.save()
    packet.seek(0)

    new_pdf = PdfReader(packet)
    writer.add_page(new_pdf.pages[0])

    output_stream = BytesIO()
    writer.write(output_stream)
    output_stream.seek(0)
    return output_stream

def append_text_to_txt(original_file, appended_text):
    original_file.seek(0)
    original_content = original_file.read().decode('utf-8')
    full_content = original_content + "\n\n=== 优化内容 ===\n" + appended_text
    return BytesIO(full_content.encode('utf-8'))

# --------- 页面主内容 ---------
tab1, tab2 = st.tabs(["📝 优化文档", "📜 优化历史记录"])

with tab1:
    st.title("🎯 报告优化助手")

    roles_prompts = {
        "施工工程师": "你是施工工程领域专家，请帮我优化以下报告内容，使其更符合施工现场实际需求和流程。",
        "设计工程师": "你是设计工程领域专家，请帮我优化以下报告内容，使其更符合设计规范和技术要求。",
        "运营专员": "你是运营管理专家，请帮我优化以下报告内容，提升运营效率和管理流程。",
        "投标专员": "你是投标专家，请帮我优化以下报告内容，突出投标优势和风险控制。",
        "无角色（原文分析）": ""
    }

    with st.expander("📁 上传文件及角色选择", expanded=True):
        uploaded_file = st.file_uploader("上传文件（支持 PDF, Word, TXT）", type=["pdf", "docx", "txt"])

        # 文件上传成功提示逻辑
        if uploaded_file is not None:
            # 只要上传了新文件且跟之前确认文件名不一样时才提示
            if ("last_uploaded_filename" not in st.session_state or
                    st.session_state.last_uploaded_filename != uploaded_file.name):
                st.session_state.last_uploaded_filename = uploaded_file.name
                st.success(f"✅ 文件上传成功：{uploaded_file.name}")


        # 角色选择，带spinner和切换成功提示
        def on_role_change_1():
            st.session_state.role_switching_1 = True
            st.session_state.role_select_1 = st.session_state.role_temp_1
            st.session_state.role_switch_info_1 = ""


        role_temp_1 = st.selectbox(
            "👥 选择角色",
            options=list(roles_prompts.keys()),
            index=list(roles_prompts.keys()).index(st.session_state.role_select_1),
            key="role_temp_1",
            on_change=on_role_change_1,
        )

        if st.session_state.role_switching_1:
            with st.spinner("正在切换角色，请稍候..."):
                time.sleep(1)
            st.session_state.role_switching_1 = False
            st.session_state.role_switch_info_1 = f"✅ 已成功切换角色 “{st.session_state.role_select_1}”"

        if st.session_state.role_switch_info_1:
            st.info(st.session_state.role_switch_info_1)

        # 后续调用角色相关时用 st.session_state.role_select_1
        role = st.session_state.role_select_1
        # 优化方向文本框，绑定session_state，保持输入内容
        user_direction = st.text_area("✔ 可选：请输入优化方向或重点（可留空）", height=80, key="user_direction")

    if st.button("🚀 开始优化"):
        role = st.session_state.role_select_1
        # 上传文件优先用刚上传的，如果没有用之前确认过的

        current_file = uploaded_file if uploaded_file is not None else st.session_state.confirmed_uploaded_file

        if current_file is None:
            st.warning("请先上传文件")
            st.stop()

        # 更新确认的上传文件到session_state，保持后续稳定
        st.session_state.confirmed_uploaded_file = current_file

        with st.spinner("⏳ 正在解析文件文本..."):
            if current_file.type == "application/pdf":
                text = extract_text_from_pdf(current_file)
            elif current_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
                text = extract_text_from_docx(current_file)
            elif current_file.type == "text/plain":
                text = extract_text_from_txt(current_file)
            else:
                st.error("暂不支持该文件类型")
                st.stop()

        if not text.strip():
            st.error("未能提取到有效文本，请确认文件内容")
            st.stop()

        base_prompt = roles_prompts.get(role, "")
        full_prompt = f"{base_prompt}\n\n请根据以下内容进行优化：\n{text}\n\n"
        if user_direction.strip():
            full_prompt += f"特别注意优化方向：{user_direction}\n\n"
        full_prompt += "请给出优化后的完整文本内容，以及优化建议。"

        with st.spinner("🎈 正在调用DeepSeek AI进行优化..."):
            response_text = call_deepseek_api(full_prompt)

        if not response_text:
            st.error("未收到有效响应，请稍后重试")
            st.stop()

        st.markdown("### ✅ 优化内容预览")
        st.text_area("优化结果", response_text, height=300)

        filename = current_file.name
        buffer = None

        try:
            current_file.seek(0)
            if current_file.type == "application/pdf":
                buffer = append_text_to_pdf(current_file, response_text)
            elif current_file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
                buffer = append_text_to_docx(current_file, response_text)
            elif current_file.type == "text/plain":
                buffer = append_text_to_txt(current_file, response_text)
        except Exception as e:
            st.error(f"生成下载文件失败: {e}")
            st.stop()

        if buffer:
            st.download_button("📥 下载优化后文件", data=buffer, file_name=f"优化_{filename}", mime=current_file.type)
            st.warning("优化内容已附加至原文档末尾", icon="📎")

            # 添加到历史记录
            st.session_state.history.append({
                "时间": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "文件名": filename,
                "角色": role,
                "优化方向": user_direction.strip() or "（无）",
                "优化内容": response_text
            })

with tab2:
    st.subheader("📜 历史优化记录")

    # 保护性初始化，防止被意外覆盖成非列表或 None
    if "history" not in st.session_state or not isinstance(st.session_state.history, list):
        st.session_state.history = []

    if not st.session_state.history:
        st.info("暂无历史记录。请先在“📝 优化文档”中执行操作。")
    else:
        for record in reversed(st.session_state.history):
            if not isinstance(record, dict):
                continue
            with st.expander(f"📄 {record.get('文件名', '未知文件')} - {record.get('时间', '未知时间')}"):
                st.markdown(f"**🕒 时间**：{record.get('时间', '')}")
                st.markdown(f"**👤 角色**：{record.get('角色', '')}")
                st.markdown(f"**📌 优化方向**：{record.get('优化方向', '')}")
                st.markdown("**📘 优化内容预览：**")
                st.text_area("优化内容", record.get('优化内容', ''), height=200, disabled=True,
                             label_visibility="collapsed")

st.markdown("""
        <style>
        .footer {
            text-align: center;
            font-size: 12px;
            color: #888888;
            margin-top: 40px;
            padding-bottom: 20px;
        }
        </style>
        <div class="footer">
            © 2025 工程过程增效 AI 工具 | Powered by DeepSeek API | Developed by Luke Yang
        </div>
    """, unsafe_allow_html=True)

with st.sidebar:
    st.markdown("""
    <style>

    section[data-testid="stSidebar"] {
        width: 350px !important;
        min-width: 150px !important;
        max-width: 550px !important;
    }
    </style>
    """, unsafe_allow_html=True)

    if "theme" not in st.session_state:
        st.session_state.theme = "🌞明亮模式"

    selected_theme = st.selectbox("🎨 主题切换", ["🌞明亮模式", "🌚暗黑模式"],
                                  index=["🌞明亮模式", "🌚暗黑模式"].index(st.session_state.theme))
    if selected_theme != st.session_state.theme:
        st.session_state.theme = selected_theme
        st.rerun()

    st.markdown("#### 用户信息")
    if "username" not in st.session_state:
        st.session_state.username = "登山者"

    user_input = st.text_input("👤 用户名", value=st.session_state.username, key="username_input")
    st.session_state.username = user_input

    st.markdown(f"""
    <div style='margin-top: -10px; margin-bottom: 15px; font-size: 16px;'>
        👋 欢迎回来，<b>{st.session_state.username}</b>！
    </div>
    """, unsafe_allow_html=True)

    st.markdown("## 🛠️ 帮助与支持")
    st.markdown("[📚 使用文档](https://api-docs.deepseek.com/)")
    st.markdown("[📈 账号流量](https://platform.deepseek.com/usage)")
    st.markdown("[💬 微信公众号](https://mp.weixin.qq.com/s/49K4_GBv9Eu9Pfr0lWadQw)")
    now = datetime.datetime.now()
    st.markdown(
        f"<div style='color: gray;margin-top: -0px; margin-bottom: -805px;'>🕙 今天是：{now.strftime('%Y-%m-%d')}</div>",
        unsafe_allow_html=True)
    st.markdown("---")

    is_dark = st.session_state.theme == "🌚暗黑模式"
    title_color = "#FFFFFF" if is_dark else "#000000"
    subtitle_color = "#AAAAAA" if is_dark else "#999999"
    st.image("图片.png", use_container_width=True)

    import base64

    with open("static/logo_font.ttf", "rb") as f:
        font_data = f.read()

    font_base64 = base64.b64encode(font_data).decode()
    st.markdown(f"""
    <style>
    @font-face {{
        font-family: 'logo_font';
        src: url(data:font/truetype;charset=utf-8;base64,{font_base64}) format('truetype');
    }}
    .custom-title {{
        font-family: 'logo_font';
        font-size: 18px;
        font-weight:regular;
        color: {title_color};
        text-align: center;
        letter-spacing: 3px;
        margin-top: -5px;
        margin-bottom: -25px;
    }}
    </style>

    <div class="custom-title">
        工程过程增效AI工具
    </div>
    """, unsafe_allow_html=True)

    st.markdown(f"""

    <div style='text-align: center; color: {subtitle_color}; font-size: 12px; margin-bottom: -40px;'>
        Based on DeepSeek | Version V-0.1
    </div>
    """, unsafe_allow_html=True)